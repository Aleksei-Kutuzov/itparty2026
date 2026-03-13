from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List
import copy


class DocMerger:
    def __init__(self):
        self.target_doc = None

    def merge_with_settings(self,
                            file_paths: List[str],
                            output_path: str,
                            add_page_breaks: bool = False,
                            insert_separator: bool = False,
                            separator_text: str = "---"):
        """Объединяет документы с полным сохранением форматирования"""

        self.target_doc = Document(file_paths[0])

        for i, file_path in enumerate(file_paths[1:], 1):
            if add_page_breaks:
                self.target_doc.add_page_break()

            if insert_separator:
                sep_para = self.target_doc.add_paragraph(separator_text)
                sep_para.style = 'Heading 3'
                sep_para.alignment = 1

            source_doc = Document(file_path)
            self._copy_content_deep(source_doc)

        self.target_doc.save(output_path)
        return self.target_doc

    def _copy_content_deep(self, source_doc: Document):
        """Глубокое копирование содержимого с сохранением форматирования"""

        # Копируем параграфы
        for source_para in source_doc.paragraphs:
            self._copy_paragraph_deep(source_para)

        # Копируем таблицы
        for source_table in source_doc.tables:
            self._copy_table_deep(source_table)

    def _copy_paragraph_deep(self, source_para):
        """Копирует параграф с полным форматированием"""
        target_para = self.target_doc.add_paragraph()

        # Копируем стиль параграфа
        if source_para.style:
            try:
                target_para.style = source_para.style
            except:
                pass

        # Копируем выравнивание
        target_para.alignment = source_para.alignment

        # Копируем все runs с форматированием
        for source_run in source_para.runs:
            target_run = target_para.add_run(source_run.text)
            self._copy_run_formatting_deep(source_run, target_run)

        # Копируем свойства параграфа (отступы, интервалы)
        self._copy_paragraph_properties(source_para, target_para)

    def _copy_run_formatting_deep(self, source_run, target_run):
        """Копирует ВСЕ свойства форматирования run"""

        # Основные атрибуты
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.strike = source_run.strike
        target_run.subscript = source_run.subscript
        target_run.superscript = source_run.superscript
        target_run.hidden = source_run.hidden
        target_run.emphasis_mark = source_run.emphasis_mark
        target_run.no_proof = source_run.no_proof
        target_run.snap_to_grid = source_run.snap_to_grid
        target_run.spell = source_run.spell

        # Шрифт
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
        if source_run.font.highlight_color:
            target_run.font.highlight_color = source_run.font.highlight_color

        # Копирование XML свойств run (для полного сохранения)
        try:
            source_rPr = source_run._element.get_or_add_rPr()
            target_rPr = target_run._element.get_or_add_rPr()

            for child in source_rPr:
                if child.tag not in [qn('w:sz'), qn('w:szCs'), qn('w:rFonts')]:
                    target_rPr.append(copy.deepcopy(child))
        except:
            pass

    def _copy_paragraph_properties(self, source_para, target_para):
        """Копирует свойства параграфа (отступы, интервалы, нумерация)"""
        try:
            source_pPr = source_para._element.get_or_add_pPr()
            target_pPr = target_para._element.get_or_add_pPr()

            # Копируем важные свойства
            for tag in ['w:pStyle', 'w:numPr', 'w:spacing', 'w:ind',
                        'w:jc', 'w:textAlignment', 'w:pBdr']:
                source_elem = source_pPr.find(qn(tag))
                if source_elem is not None:
                    target_pPr.append(copy.deepcopy(source_elem))
        except:
            pass

    def _copy_table_deep(self, source_table):
        """Копирует таблицу с полным форматированием"""

        # Создаём таблицу
        target_table = self.target_doc.add_table(
            rows=len(source_table.rows),
            cols=len(source_table.columns)
        )

        # Копируем стиль таблицы
        if source_table.style:
            try:
                target_table.style = source_table.style
            except:
                pass

        # Копируем свойства таблицы (границы, выравнивание)
        self._copy_table_properties(source_table, target_table)

        # Копируем ширину столбцов
        for i, source_col in enumerate(source_table.columns):
            if i < len(target_table.columns):
                try:
                    target_table.columns[i].width = source_col.width
                except:
                    pass

        # Копируем содержимое ячеек
        for i, source_row in enumerate(source_table.rows):
            target_row = target_table.rows[i]

            for j, source_cell in enumerate(source_row.cells):
                target_cell = target_row.cells[j]

                # Копируем ширину ячейки
                try:
                    target_cell.width = source_cell.width
                except:
                    pass

                # Копируем форматирование ячейки
                self._copy_cell_formatting(source_cell, target_cell)

                # Копируем параграфы в ячейке
                target_cell.paragraphs[0].clear()

                for k, source_para in enumerate(source_cell.paragraphs):
                    if k == 0:
                        self._copy_paragraph_deep_to_cell(
                            source_para,
                            target_cell.paragraphs[0]
                        )
                    else:
                        new_para = target_cell.add_paragraph()
                        self._copy_paragraph_deep_to_cell(
                            source_para,
                            new_para
                        )

    def _copy_table_properties(self, source_table, target_table):
        """Копирует свойства таблицы"""
        try:
            source_tblPr = source_table._element.tblPr
            target_tblPr = target_table._element.get_or_add_tblPr()

            # Копируем границы, выравнивание, shading
            for tag in ['w:tblBorders', 'w:tblStyle', 'w:jc',
                        'w:tblCellMar', 'w:tblLayout']:
                source_elem = source_tblPr.find(qn(tag))
                if source_elem is not None:
                    target_tblPr.append(copy.deepcopy(source_elem))
        except:
            pass

    def _copy_cell_formatting(self, source_cell, target_cell):
        """Копирует форматирование ячейки таблицы"""
        try:
            source_tcPr = source_cell._element.tcPr
            target_tcPr = target_cell._element.get_or_add_tcPr()

            # Копируем свойства ячейки
            for tag in ['w:tcBorders', 'w:shd', 'w:vAlign',
                        'w:tcW', 'w:gridSpan']:
                source_elem = source_tcPr.find(qn(tag))
                if source_elem is not None:
                    target_tcPr.append(copy.deepcopy(source_elem))
        except:
            pass

    def _copy_paragraph_deep_to_cell(self, source_para, target_para):
        """Копирует параграф в ячейку таблицы"""
        if source_para.style:
            try:
                target_para.style = source_para.style
            except:
                pass

        target_para.alignment = source_para.alignment

        for source_run in source_para.runs:
            target_run = target_para.add_run(source_run.text)
            self._copy_run_formatting_deep(source_run, target_run)

        self._copy_paragraph_properties(source_para, target_para)


from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy


class XMLDocMerger:
    def __init__(self):
        self.target_doc = None

    def merge(self, file_paths, output_path, force_table_borders=True):
        self.target_doc = Document(file_paths[0])

        self._copy_table_styles(file_paths)

        for file_path in file_paths[1:]:
            source_doc = Document(file_path)

            for element in source_doc.element.body:
                if element.tag.endswith('p'):
                    self._copy_paragraph_xml(element)
                elif element.tag.endswith('tbl'):
                    self._copy_table_xml(element, force_borders=force_table_borders)

        self.target_doc.save(output_path)

    def _copy_table_styles(self, file_paths):
        for file_path in file_paths[1:]:
            source_doc = Document(file_path)
            source_styles = source_doc.styles

            for style in source_styles:
                if style.name not in self.target_doc.styles:
                    try:
                        if 'table' in str(style.type).lower():
                            new_style = self.target_doc.styles.add_style(
                                style.name,
                                style.type
                            )
                    except:
                        pass

    def _copy_paragraph_xml(self, source_element):
        target_element = copy.deepcopy(source_element)
        self.target_doc.element.body.append(target_element)

    def _copy_table_xml(self, source_element, force_borders=True):
        target_element = copy.deepcopy(source_element)

        if force_borders:
            self._ensure_table_borders(target_element)

        self.target_doc.element.body.append(target_element)

    def _ensure_table_borders(self, table_element):
        tblPr = table_element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table_element.insert(0, tblPr)


        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)


        border_config = {
            'top': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
            'left': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
            'bottom': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
            'right': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
            'insideH': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'},
            'insideV': {'sz': '4', 'val': 'single', 'color': '000000', 'space': '0'}
        }


        for border_name, config in border_config.items():
            existing_border = tblBorders.find(qn(f'w:{border_name}'))
            if existing_border is None:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:sz'), config['sz'])
                border_elem.set(qn('w:val'), config['val'])
                border_elem.set(qn('w:color'), config['color'])
                border_elem.set(qn('w:space'), config['space'])
                tblBorders.append(border_elem)
            else:

                for attr, value in config.items():
                    existing_border.set(qn(f'w:{attr}'), value)


        for row in table_element.findall(qn('w:tr')):
            for cell in row.findall(qn('w:tc')):
                self._ensure_cell_borders(cell)

    def _ensure_cell_borders(self, cell_element):
        """Добавляет границы ячейкам"""
        tcPr = cell_element.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr')

            if len(cell_element) > 0:
                cell_element.insert(1, tcPr)
            else:
                cell_element.append(tcPr)

        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)


        for border_name in ['top', 'left', 'bottom', 'right']:
            existing_border = tcBorders.find(qn(f'w:{border_name}'))
            if existing_border is None:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:sz'), '4')
                border_elem.set(qn('w:val'), 'single')
                border_elem.set(qn('w:color'), '000000')
                border_elem.set(qn('w:space'), '0')
                tcBorders.append(border_elem)
